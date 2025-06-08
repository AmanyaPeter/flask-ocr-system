# app/utils.py
import os
import cv2
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import letter
from pytesseract import Output

# If Tesseract is not in your PATH, include the following line
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def preprocess_image(image):
    """Preprocesses an image for better OCR results."""
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # Apply thresholding
    thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    # Optional: Apply denoising
    denoised = cv2.fastNlMeansDenoising(thresh, h=30)
    return denoised

def get_ocr_data(image, lang='eng'):
    """Performs OCR and returns text and confidence data."""
    # Use image_to_data to get detailed information
    data = pytesseract.image_to_data(image, lang=lang, output_type=Output.DICT)
    return data

def process_file(filepath, lang='eng'):
    """
    Processes a single file (image or PDF) and returns OCR results.
    Returns a dictionary containing text, page images, and confidence data.
    """
    filename = os.path.basename(filepath)
    results = {'filename': filename, 'pages': []}
    
    if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
        image_cv = cv2.imread(filepath)
        processed_image = preprocess_image(image_cv)
        ocr_data = get_ocr_data(processed_image, lang)
        
        # Save a preview image
        preview_path = os.path.join('processed', f"preview_{filename}")
        cv2.imwrite(os.path.join('static', preview_path), cv2.imread(filepath))

        results['pages'].append({
            'page_num': 1,
            'text': " ".join(ocr_data['text']).strip(),
            'ocr_data': ocr_data,
            'preview_image': preview_path
        })
        results['page_count'] = 1

    elif filename.lower().endswith('.pdf'):
        images_from_pdf = convert_from_path(filepath)
        results['page_count'] = len(images_from_pdf)
        
        for i, pil_image in enumerate(images_from_pdf):
            page_num = i + 1
            temp_image_path = f"temp_page_{page_num}_{filename}.png"
            pil_image.save(temp_image_path, 'PNG')

            image_cv = cv2.imread(temp_image_path)
            processed_image = preprocess_image(image_cv)
            ocr_data = get_ocr_data(processed_image, lang)
            
            # Save a preview image
            preview_filename = f"preview_page_{page_num}_{filename}"
            preview_path = os.path.join('processed', preview_filename)
            pil_image.save(os.path.join('static', preview_path), 'PNG')
            
            results['pages'].append({
                'page_num': page_num,
                'text': " ".join(ocr_data['text']).strip(),
                'ocr_data': ocr_data,
                'preview_image': preview_path
            })
            os.remove(temp_image_path) # Clean up temp file

    return results

def create_downloadable_file(results, file_format):
    """Creates a downloadable file (txt, docx, or searchable pdf)."""
    full_text = "\n\n".join([page['text'] for page in results['pages']])
    
    if file_format == 'txt':
        return full_text, 'text/plain'
        
    elif file_format == 'docx':
        doc = Document()
        doc.add_heading(f"OCR Result for {results['filename']}", 0)
        doc.add_paragraph(full_text)
        
        # Save to a temporary in-memory buffer
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    elif file_format == 'pdf':
         # Use pytesseract's built-in functionality for searchable PDF
        pdf_bytes = b''
        for page in results['pages']:
             original_image_path = os.path.join('static', page['preview_image'])
             pdf_bytes += pytesseract.image_to_pdf_or_hocr(original_image_path, extension='pdf', lang='eng') # Use same lang as OCR
        
        return pdf_bytes, 'application/pdf'
    
    return None, None
